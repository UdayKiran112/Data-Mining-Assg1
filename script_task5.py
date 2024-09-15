from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import pandas as pd
import numpy as np
from sklearn.preprocessing import StandardScaler
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.decomposition import PCA

# Create a new Document
doc = Document()
doc.add_heading('Dimensionality Reduction Analysis Report', level=1)

# Introduction Section
doc.add_heading('1. Introduction', level=2)
doc.add_paragraph(
    "This document outlines the approach taken to analyze pairwise similarity among data objects after applying dimensionality reduction techniques using Principal Component Analysis (PCA). "
    "The task involves repeating the pairwise similarity analysis from Task 4, but with a reduced-dimensional dataset, and providing insights on the effectiveness of dimensionality reduction."
)

# Dimensionality Reduction Section
doc.add_heading('2. Dimensionality Reduction', level=2)
doc.add_heading('2.1 Dataset:', level=3)
doc.add_paragraph(
    "The dataset used for this analysis is the 'Air Quality UCI' dataset, which includes various air quality metrics."
)
doc.add_heading('2.2 Dimensionality Reduction Technique:', level=3)
doc.add_paragraph(
    "Principal Component Analysis (PCA) was applied to reduce the number of features in the dataset. PCA is a technique that transforms the data into a set of orthogonal components, capturing the most variance with fewer dimensions."
)
doc.add_heading('2.3 PCA Implementation:', level=3)
doc.add_paragraph(
    "Number of Components: 5\n"
    "Preprocessing Steps:\n"
    "- Replaced missing values (-200) with NaN.\n"
    "- Filled missing values in numeric columns with the median of the respective columns.\n"
    "- Standardized numeric features."
)

# Add Python Code Section
doc.add_heading('3. Python Code', level=2)
code = """
import pandas as pd
import numpy as np
from sklearn.preprocessing import StandardScaler
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.decomposition import PCA

# Load and preprocess the dataset
df = pd.read_excel('AirQualityUCI.xlsx')

# Replace -200 with NaN
df.replace(-200, np.nan, inplace=True)

# Separate numeric columns
numeric_df = df.select_dtypes(include=[np.number])

# Compute medians for numeric columns and fill missing values
df[numeric_df.columns] = numeric_df.fillna(numeric_df.median())

# Reset index and sample 20 data objects
df.reset_index(drop=True, inplace=True)
sampled_df = df.sample(n=20, random_state=1)

# Extract numeric features for similarity calculation
features = sampled_df.select_dtypes(include=[np.number])

# Standardize features
scaler = StandardScaler()
scaled_features = scaler.fit_transform(features)

# Compute Cosine Similarity with original features
cosine_sim_original = cosine_similarity(scaled_features)

def get_max_similarity_pair(similarity_matrix):
    np.fill_diagonal(similarity_matrix, 0)  # Set diagonal to 0 to avoid self-similarity
    max_similarity_idx = np.unravel_index(np.argmax(similarity_matrix, axis=None), similarity_matrix.shape)
    max_similarity_score = similarity_matrix[max_similarity_idx]
    return max_similarity_idx, max_similarity_score

def print_pair_info(df, idx, measure_name, score):
    # Extract numeric columns for displaying pairs
    numeric_columns = df.select_dtypes(include=[np.number]).columns
    pair_1 = df.iloc[idx[0]][numeric_columns]
    pair_2 = df.iloc[idx[1]][numeric_columns]
    
    print(f"\\n{measure_name} Similarity:")
    print(f"Pair with maximum similarity (Index {idx[0]} and {idx[1]}):")
    print(f"Pair 1:\\n{pair_1}\\n")
    print(f"Pair 2:\\n{pair_2}\\n")
    print(f"Similarity Score: {score:.4f}")

    # Check if they are really similar
    if score > 0.8:  # You can adjust this threshold based on your context
        print("The pairs are really similar.")
    else:
        print("The pairs are not very similar.")

# Get the pair with maximum similarity and the similarity score in the original dataset
original_max_idx, original_max_score = get_max_similarity_pair(cosine_sim_original)
print_pair_info(sampled_df, original_max_idx, 'Cosine Similarity (Original)', original_max_score)

# Apply PCA for dimensionality reduction
pca = PCA(n_components=5)  # Reduce to 5 dimensions
pca_features = pca.fit_transform(scaled_features)

# Compute Cosine Similarity with PCA-reduced features
cosine_sim_pca = cosine_similarity(pca_features)

# Get the pair with maximum similarity and the similarity score after PCA
pca_max_idx, pca_max_score = get_max_similarity_pair(cosine_sim_pca)
print_pair_info(sampled_df, pca_max_idx, 'Cosine Similarity (PCA)', pca_max_score)
"""
doc.add_paragraph(code, style='Normal')

# Results Section
doc.add_heading('4. Results', level=2)

doc.add_heading('4.1 Cosine Similarity (Original Dataset):', level=3)
doc.add_paragraph(
    "Pair with maximum similarity (Index 0 and 2):\n"
    "Pair 1:\n"
    "CO(GT)                   1.1\n"
    "PT08.S1(CO)      1047.333333\n"
    "NMHC(GT)                74.0\n"
    "C6H6(GT)            4.932008\n"
    "PT08.S2(NMHC)          760.0\n"
    "NOx(GT)                 64.0\n"
    "PT08.S3(NOx)          1032.0\n"
    "NO2(GT)                 74.0\n"
    "PT08.S4(NO2)     1378.666667\n"
    "PT08.S5(O3)           1003.0\n"
    "T                  11.466667\n"
    "RH                 61.433333\n"
    "AH                  0.830289\n"
    "\n"
    "Pair 2:\n"
    "CO(GT)                 1.8\n"
    "PT08.S1(CO)         1129.5\n"
    "NMHC(GT)              56.0\n"
    "C6H6(GT)          5.191654\n"
    "PT08.S2(NMHC)        773.0\n"
    "NOx(GT)               70.0\n"
    "PT08.S3(NOx)       1130.25\n"
    "NO2(GT)               82.0\n"
    "PT08.S4(NO2)       1451.75\n"
    "PT08.S5(O3)         1050.5\n"
    "T                     12.1\n"
    "RH               61.100001\n"
    "AH                0.860316\n"
    "\n"
    "Similarity Score: 0.9681\n"
    "Insight: The pairs are highly similar, showing strong consistency in the high-dimensional feature space."
)

doc.add_heading('4.2 Cosine Similarity (PCA-Reduced Dataset):', level=3)
doc.add_paragraph(
    "Pair with maximum similarity (Index 10 and 16):\n"
    "Pair 1:\n"
    "CO(GT)                 1.8\n"
    "PT08.S1(CO)         1125.0\n"
    "NMHC(GT)             150.0\n"
    "C6H6(GT)         12.170581\n"
    "PT08.S2(NMHC)       1055.5\n"
    "NOx(GT)              179.8\n"
    "PT08.S3(NOx)         645.0\n"
    "NO2(GT)              109.0\n"
    "PT08.S4(NO2)        1715.5\n"
    "PT08.S5(O3)         929.25\n"
    "T                33.825001\n"
    "RH                  29.225\n"
    "AH                1.514256\n"
    "\n"
    "Pair 2:\n"
    "CO(GT)                 2.1\n"
    "PT08.S1(CO)        1180.75\n"
    "NMHC(GT)             150.0\n"
    "C6H6(GT)         12.816446\n"
    "PT08.S2(NMHC)       1077.5\n"
    "NOx(GT)              159.0\n"
    "PT08.S3(NOx)         706.0\n"
    "NO2(GT)              122.0\n"
    "PT08.S4(NO2)       1869.75\n"
    "PT08.S5(O3)        1139.75\n"
    "T                     35.4\n"
    "RH                  27.825\n"
    "AH                1.573265\n"
    "\n"
    "Similarity Score: 0.9854\n"
    "Insight: The similarity score is slightly higher after PCA, indicating that PCA has preserved or enhanced the similarity between the pairs. The pairs are still highly similar, suggesting that PCA effectively captures the essential relationships within the data."
)

# Dimension Details After PCA
doc.add_heading('4.3 Dimension Details After PCA', level=3)
doc.add_paragraph(
    "Explained Variance Ratio for each Principal Component:\n"
    "PC1: 0.5914\n"
    "PC2: 0.1807\n"
    "PC3: 0.1031\n"
    "PC4: 0.0756\n"
    "PC5: 0.0274\n"
)
doc.add_paragraph(
    "Cumulative Explained Variance Ratio:\n"
    "PC1: 0.5914\n"
    "PC2: 0.7721\n"
    "PC3: 0.8751\n"
    "PC4: 0.9508\n"
    "PC5: 0.9781\n"
)
doc.add_paragraph(
    "Principal Components (Loadings):\n"
    "PC1: [ 0.34492862  0.33707433  0.06322231  0.35426103  0.35201093  0.29817699\n"
    "       -0.32442131  0.32720973  0.27648803  0.34063724  0.11016456 -0.07342537\n"
    "        0.05257104]\n"
    "PC2: [ 0.14271999  0.15355303 -0.23922705  0.01501506 -0.04839629  0.2112071\n"
    "        0.04776545  0.09061993 -0.30407794  0.14826984 -0.61121721  0.33437171\n"
    "       -0.49123827]\n"
    "PC3: [-0.07794254  0.08073365 -0.5495173  -0.09925671 -0.02905356 -0.09628982\n"
    "       -0.00213746 -0.03818745  0.29207126  0.16388612  0.02060397  0.58232164\n"
    "        0.46078183]\n"
    "PC4: [ 0.0972446   0.13124127 -0.63018628  0.10121408  0.14239622 -0.32592841\n"
    "        0.1206921  -0.22128964  0.16736476  0.0325719   0.09674064 -0.47163934\n"
    "       -0.33915212]\n"
    "PC5: [ 0.09345098 -0.26301915 -0.21866943 -0.01335577 -0.17048998  0.45368084\n"
    "        0.61369325  0.44428563  0.11527447 -0.01945856  0.08643975 -0.18914318\n"
    "        0.0909784 ]\n"
    "\n"
    "Note: The principal components are the directions in which the data varies the most. The loadings indicate the contribution of each original feature to the principal components."
)

# Feature Variance Details
doc.add_heading('4.4 Feature Variance Before and After Scaling', level=3)
doc.add_paragraph(
    "Feature Variance Before Scaling:\n"
    "Feature 1: 2.0929\n"
    "Feature 2: 46644.5713\n"
    "Feature 3: 693.0000\n"
    "Feature 4: 53.0124\n"
    "Feature 5: 81560.1429\n"
    "Feature 6: 42995.0089\n"
    "Feature 7: 67274.7888\n"
    "Feature 8: 2055.3694\n"
    "Feature 9: 114807.4505\n"
    "Feature 10: 175210.1574\n"
    "Feature 11: 69.1173\n"
    "Feature 12: 230.2564\n"
    "Feature 13: 0.1494\n"
)
doc.add_paragraph(
    "Feature Variance After Scaling:\n"
    "Feature 1: 1.0526\n"
    "Feature 2: 1.0526\n"
    "Feature 3: 1.0526\n"
    "Feature 4: 1.0526\n"
    "Feature 5: 1.0526\n"
    "Feature 6: 1.0526\n"
    "Feature 7: 1.0526\n"
    "Feature 8: 1.0526\n"
    "Feature 9: 1.0526\n"
    "Feature 10: 1.0526\n"
    "Feature 11: 1.0526\n"
    "Feature 12: 1.0526\n"
    "Feature 13: 1.0526\n"
)

# Conclusion Section
doc.add_heading('5. Conclusion', level=2)
doc.add_paragraph(
    "The analysis demonstrates that PCA effectively reduces dimensionality while retaining the critical relationships within the dataset. Both the original and PCA-reduced datasets show high similarity scores between the most similar pairs. "
    "The slight increase in similarity score post-PCA suggests that dimensionality reduction can enhance data representation by removing noise and focusing on the principal components that capture the most variance."
)

# Save the document
doc.save('Dimensionality_Reduction_Analysis_Report.docx')
