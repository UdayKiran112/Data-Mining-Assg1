from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement

# Create a new Document
doc = Document()

# Title
doc.add_heading('Data Pre-processing', level=1)

# Task Description
doc.add_heading('Task 3: Normalization of Selected Attribute', level=2)

# Attribute Details
doc.add_heading('Attribute Selected: PT08.S1(CO)', level=3)
doc.add_paragraph(
    "The selected attribute for normalization is `PT08.S1(CO)`, which represents the concentration "
    "of Carbon Monoxide (CO) measured in the air. This attribute is part of the Air Quality UCI dataset, "
    "which is used to monitor air pollution levels. The attribute values are numerical, and it is essential "
    "to preprocess these values to ensure they are in a consistent scale for further analysis and machine learning tasks."
)

# Normalization Technique
doc.add_heading('Normalization Technique', level=3)
doc.add_paragraph(
    "Normalization is a technique used to scale the data to a fixed range, typically [0, 1]. This is particularly "
    "useful when dealing with numerical data that varies significantly in magnitude. In this case, we use Min-Max "
    "Scaling (or Min-Max Normalization) which transforms the data to fit within a specified range."
)

doc.add_paragraph(
    "Why Min-Max Scaling?\n"
    "- **Uniform Scaling**: It scales the data between 0 and 1, which is helpful for many machine learning algorithms "
    "that are sensitive to the scale of input features.\n"
    "- **Preservation of Relationships**: This technique preserves the relationships between data points, making it suitable "
    "for attributes where the original scale is not relevant for analysis or modeling."
)

# Python Code
doc.add_heading('Python Code for Normalization', level=3)
code = """\
import pandas as pd
from sklearn.preprocessing import MinMaxScaler

# Load the dataset
df = pd.read_excel('AirQualityUCI.xlsx')

# Select an attribute for normalization
attribute = 'PT08.S1(CO)'

# Handle missing values (-200 values treated as NaN)
df[attribute] = df[attribute].replace(-200, pd.NA)

# Fill missing values in the selected attribute with its median
median_value = df[attribute].median()
df[attribute] = df[attribute].fillna(median_value)

# Normalize the attribute
scaler = MinMaxScaler()
df[attribute + '_normalized'] = scaler.fit_transform(df[[attribute]])

# Print results
print(f"Normalized values of {attribute}:")
print(df[[attribute, attribute + '_normalized']].head())
"""
# Add code as preformatted text
p = doc.add_paragraph()
p.add_run(code).font.name = 'Courier New'
p.style.font.size = Pt(10)

# Explanation of the Code
doc.add_heading('Explanation of the Code', level=3)
doc.add_paragraph(
    "1. **Loading the Dataset**: The dataset is loaded from an Excel file.\n"
    "2. **Handling Missing Values**: We replace erroneous `-200` values with `NaN` to identify and handle missing values correctly.\n"
    "3. **Filling Missing Values**: Missing values are filled using the median of the attribute to ensure that the data is complete and consistent.\n"
    "4. **Normalization**: Min-Max Scaling is applied to the attribute to normalize its values between 0 and 1. This ensures that all values are on a uniform scale."
)

# Save the document
doc.save('Data_Preprocessing_Task3.docx')

print("Document created successfully.")
