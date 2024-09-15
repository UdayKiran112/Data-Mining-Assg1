from docx import Document
from docx.shared import Pt

# Create a new Document
doc = Document()

# Title
doc.add_heading('Task 4 - Pairwise Similarity Analysis', level=1)

# Introduction
doc.add_heading('1. Introduction', level=2)
doc.add_paragraph(
    "This document describes the approach taken to analyze pairwise similarity among data objects selected from the dataset. "
    "The task involves sampling a subset of data, computing similarity scores using appropriate measures, identifying the pair with the maximum similarity, "
    "and evaluating if the selected pair is truly similar."
)

# Sampling Strategy
doc.add_heading('2. Sampling Strategy', level=2)
doc.add_paragraph(
    "To address the task, the following sampling strategy was employed:"
)
doc.add_paragraph(
    "1. Dataset: The dataset used for this analysis is the 'Air Quality UCI' dataset, which contains various air quality metrics."
)
doc.add_paragraph(
    "2. Sampling Technique: A random sampling approach was used to select 20 data objects from the dataset. This random sampling ensures that the selected subset is representative of the overall data distribution."
)

# Similarity Measure
doc.add_heading('3. Similarity Measure', level=2)
doc.add_paragraph(
    "The similarity between data objects was computed using the Cosine Similarity measure. We used cosine similarity as the dataset is high-dimensional with 15 features. "
    "This measure evaluates the cosine of the angle between two non-zero vectors, providing a value between -1 and 1. A value closer to 1 indicates higher similarity."
)

# Preprocessing Code
doc.add_heading('3.1. Preprocessing', level=3)
doc.add_paragraph(
    "The following Python code was used for preprocessing:"
)
code = """\
import pandas as pd
import numpy as np
from sklearn.preprocessing import StandardScaler

# Load and preprocess your dataset
df = pd.read_excel('AirQualityUCI.xlsx')

# Replace -200 with NaN
df.replace(-200, np.nan, inplace=True)

# Separate numeric and non-numeric columns
numeric_df = df.select_dtypes(include=[np.number])
non_numeric_df = df.select_dtypes(exclude=[np.number])

# Compute medians for numeric columns and fill missing values
df[numeric_df.columns] = numeric_df.fillna(numeric_df.median())

# Reset index and sample 20 data objects
df.reset_index(drop=True, inplace=True)
sampled_df = df.sample(n=20, random_state=1)

# Extract numeric features for similarity calculation
features = sampled_df.select_dtypes(include=[np.number])

# Standardize features
scaler = StandardScaler()
scaled_features = scaler.fit_transform(features)
"""
doc.add_paragraph(code)

# Cosine Similarity Computation Code
doc.add_heading('3.2. Cosine Similarity Computation', level=3)
doc.add_paragraph(
    "The following Python code was used to compute cosine similarity:"
)
code = """\
from sklearn.metrics.pairwise import cosine_similarity

# Compute Cosine Similarity
cosine_sim = cosine_similarity(scaled_features)
"""
doc.add_paragraph(code)

# Identification of Maximum Similarity Pair Code
doc.add_heading('3.3. Identification of Maximum Similarity Pair', level=3)
doc.add_paragraph(
    "The following Python code was used to identify the pair with maximum similarity and to print the results:"
)
code = """\
def get_max_similarity_pair(similarity_matrix):
    np.fill_diagonal(similarity_matrix, 0)  # Set diagonal to 0 to avoid self-similarity
    max_similarity_idx = np.unravel_index(np.argmax(similarity_matrix, axis=None), similarity_matrix.shape)
    max_similarity_score = similarity_matrix[max_similarity_idx]
    return max_similarity_idx, max_similarity_score

def print_pair_info(df, idx, measure_name, score):
    # Extract numeric columns for displaying pairs
    numeric_columns = df.select_dtypes(include=[np.number]).columns
    pair_1 = df.iloc[idx[0]][numeric_columns]
    pair_2 = df.iloc[idx[1]][numeric_columns]
    
    print(f"\\n{measure_name} Similarity:")
    print(f"Pair with maximum similarity (Index {idx[0]} and {idx[1]}):")
    print(f"Pair 1:\\n{pair_1}\\n")
    print(f"Pair 2:\\n{pair_2}\\n")
    print(f"Similarity Score: {score:.4f}")

    # Check if they are really similar
    if score > 0.8:  # You can adjust this threshold based on your context
        print("The pairs are really similar.")
    else:
        print("The pairs are not very similar.")

# Get the pair with maximum similarity and the similarity score
cosine_max_idx, cosine_max_score = get_max_similarity_pair(cosine_sim)
print_pair_info(sampled_df, cosine_max_idx, 'Cosine Similarity', cosine_max_score)
"""
doc.add_paragraph(code)

# Results
doc.add_heading('4. Results', level=2)
doc.add_paragraph(
    "The pair with the maximum similarity was identified as follows:"
)
doc.add_paragraph(
    "Cosine Similarity Similarity:"
)
doc.add_paragraph(
    "Pair with Maximum Similarity (Index 0 and 2):"
)
doc.add_paragraph(
    "Pair 1:\n"
    "CO(GT)                   1.1\n"
    "PT08.S1(CO)      1047.333333\n"
    "NMHC(GT)                74.0\n"
    "C6H6(GT)            4.932008\n"
    "PT08.S2(NMHC)          760.0\n"
    "NOx(GT)                 64.0\n"
    "PT08.S3(NOx)          1032.0\n"
    "NO2(GT)                 74.0\n"
    "PT08.S4(NO2)     1378.666667\n"
    "PT08.S5(O3)           1003.0\n"
    "T                  11.466667\n"
    "RH                 61.433333\n"
    "AH                  0.830289\n"
    "Name: 822, dtype: object\n"
)
doc.add_paragraph(
    "Pair 2:\n"
    "CO(GT)                 1.8\n"
    "PT08.S1(CO)         1129.5\n"
    "NMHC(GT)              56.0\n"
    "C6H6(GT)          5.191654\n"
    "PT08.S2(NMHC)        773.0\n"
    "NOx(GT)               70.0\n"
    "PT08.S3(NOx)       1130.25\n"
    "NO2(GT)               82.0\n"
    "PT08.S4(NO2)       1451.75\n"
    "PT08.S5(O3)         1050.5\n"
    "T                     12.1\n"
    "RH               61.100001\n"
    "AH                0.860316\n"
    "Name: 82, dtype: object\n"
)
doc.add_paragraph(
    "Similarity Score: 0.9681\n"
)
doc.add_paragraph(
    "The pairs are really similar."
)

# Conclusion
doc.add_heading('5. Conclusion', level=2)
doc.add_paragraph(
    "This document summarizes the methodology and results of the pairwise similarity analysis. The sampling strategy, similarity measure used, and the computation process were outlined, along with the evaluation of the most similar pair from the dataset."
)

# Save the document
doc.save('Task_4_Pairwise_Similarity_Analysis.docx')
